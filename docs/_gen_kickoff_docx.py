from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(3.0)
section.bottom_margin = Cm(3.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_p(text, italic=False, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.italic = italic
    run.bold = bold
    return para


def bullet(text):
    para = doc.add_paragraph(text, style="List Bullet")
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("KICKOFF DEL PROYECTO")
r.bold = True
r.font.name = "Arial"
r.font.size = Pt(14)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run(
    "Plantilla corporativa — Proyecto Integrador (Cibertec)\n"
    "Completa los campos [RELLENAR]. Los textos en cursiva son ejemplos."
)
rs.italic = True
rs.font.name = "Arial"
rs.font.size = Pt(11)

h("1. Datos generales", 1)
table = doc.add_table(rows=7, cols=2)
table.style = "Table Grid"
rows = [
    ("Nombre del sistema", "[RELLENAR]  Ej: LibreriaRent / BookLease Perú"),
    ("Nombre corto / código", "[RELLENAR]  Ej: LR-2026"),
    ("Fecha de kickoff", "[RELLENAR]  Ej: 2026-07-26"),
    ("Ciclo / aula", "[RELLENAR]  Ej: Sexto ciclo — aula XXX"),
    ("Coordinador del grupo", "[RELLENAR]"),
    ("Integrantes", "[RELLENAR]  Ej: Nombre1, Nombre2"),
    (
        "Cursos que abarca",
        "Móviles II (4696) + Proyecto Integrador (2423) + Pruebas de Software (2424)",
    ),
]
for i, (a, b) in enumerate(rows):
    table.rows[i].cells[0].text = a
    table.rows[i].cells[1].text = b

h("2. Problema en una frase", 1)
add_p("[RELLENAR]")
add_p(
    "Ejemplo: Las bibliotecas o negocios de alquiler de libros físicos pierden "
    "control de ejemplares, garantías y plazos porque el proceso se hace a mano "
    "o sin una app unificada.",
    italic=True,
)

h("3. Solución en una frase", 1)
add_p("[RELLENAR]")
add_p(
    "Ejemplo: App iOS + API Spring Boot para reservar/alquilar ejemplares con UID, "
    "pagar con Yape (alquiler + garantía), recoger y devolver en local, con login "
    "email/Facebook y asistente Watson.",
    italic=True,
)

h("4. Roles Scrum", 1)
t2 = doc.add_table(rows=4, cols=3)
t2.style = "Table Grid"
t2.rows[0].cells[0].text = "Rol"
t2.rows[0].cells[1].text = "Responsable"
t2.rows[0].cells[2].text = "Notas"
t2.rows[1].cells[0].text = "Product Owner"
t2.rows[1].cells[1].text = "[RELLENAR]"
t2.rows[1].cells[2].text = "Define prioridad del backlog"
t2.rows[2].cells[0].text = "Scrum Master"
t2.rows[2].cells[1].text = "[RELLENAR]"
t2.rows[2].cells[2].text = "Cuida el proceso Scrum / Trello"
t2.rows[3].cells[0].text = "Development Team"
t2.rows[3].cells[1].text = "[RELLENAR]"
t2.rows[3].cells[2].text = "API, iOS, pruebas, docs"
add_p("Si trabajas solo: pon tu nombre en los 3 roles y anótalo.", italic=True)

h("5. Tablero Trello", 1)
t3 = doc.add_table(rows=3, cols=2)
t3.style = "Table Grid"
t3.rows[0].cells[0].text = "URL del tablero"
t3.rows[0].cells[1].text = "[RELLENAR]  Ej: https://trello.com/b/xxxx/libreriarent"
t3.rows[1].cells[0].text = "Columnas"
t3.rows[1].cells[1].text = "Backlog | Sprint actual | En progreso | En prueba | Hecho"
t3.rows[2].cells[0].text = "Convención"
t3.rows[2].cells[1].text = (
    "Una card = una historia o tarea. Etiquetas: iOS, API, Docs, Watson, Pruebas, Bloqueado"
)

h("6. Alcance acordado (MVP)", 1)
add_p("Incluye:", bold=True)
for item in [
    "Títulos + ejemplares con UID",
    "Login email/contraseña + Facebook (Móviles U4)",
    "Pago Yape (celular + código de verificación) = alquiler + garantía",
    "Recojo y devolución solo en biblioteca",
    "Plazo del alquiler corre desde el pago",
    "REST (Spring Boot) + Core Data + Firebase",
    "Watson Assistant",
    "Roles admin y usuario",
]:
    bullet(item)

add_p("No incluye (MVP):", bold=True)
for item in [
    "Delivery a domicilio",
    "Venta definitiva de libros",
    "Multas automáticas",
    "Varias sucursales",
]:
    bullet(item)

h("7. Stack técnico acordado", 1)
t4 = doc.add_table(rows=8, cols=2)
t4.style = "Table Grid"
stack = [
    ("Capa", "Tecnología"),
    ("App móvil", "iOS, UIKit, Storyboard, MVC modular"),
    ("Persistencia local", "Core Data"),
    ("Backend", "Java Spring Boot (REST)"),
    ("Nube / auth extra", "Firebase + Facebook Login"),
    ("Gestión", "Scrum + Trello"),
    ("IA", "IBM Watson"),
    ("Pruebas", "Java (JUnit, Rest Assured, etc.) sobre la API"),
]
for i, (a, b) in enumerate(stack):
    t4.rows[i].cells[0].text = a
    t4.rows[i].cells[1].text = b

h("8. Próximo hito", 1)
t5 = doc.add_table(rows=4, cols=3)
t5.style = "Table Grid"
t5.rows[0].cells[0].text = "Hito"
t5.rows[0].cells[1].text = "Entregable"
t5.rows[0].cells[2].text = "Fecha objetivo"
t5.rows[1].cells[0].text = "AT01"
t5.rows[1].cells[1].text = (
    "Problemática, justificación, SEPTE, alcance, riesgos, viabilidad, Product Backlog"
)
t5.rows[1].cells[2].text = "[RELLENAR]"
t5.rows[2].cells[0].text = "AT02"
t5.rows[2].cells[1].text = "Release Plan, Sprint Backlogs, Plan de pruebas, Plan CI"
t5.rows[2].cells[2].text = "[RELLENAR]"
t5.rows[3].cells[0].text = "Sprint 1–2"
t5.rows[3].cells[1].text = "Incremento funcional (API + iOS)"
t5.rows[3].cells[2].text = "[RELLENAR]"

h("9. Definition of Done (DoD) básica", 1)
add_p("Una historia está Hecho cuando:")
bullet("Código en el repo / build OK")
bullet("Cumple criterio de aceptación")
bullet("Probada (manual o automatizada según el plan)")
bullet("Card movida a Hecho en Trello")
bullet("Documentado lo mínimo (si aplica)")

h("10. Acuerdos del equipo", 1)
bullet("[RELLENAR]  Ej: Daily asíncrono por WhatsApp 9:00")
bullet("[RELLENAR]  Ej: Nadie sube avance sin actualizar Trello")
bullet("[RELLENAR]  Ej: Decisiones de alcance se anotan en este doc o en AT01")

h("11. Firmas / conformidad", 1)
t6 = doc.add_table(rows=2, cols=4)
t6.style = "Table Grid"
t6.rows[0].cells[0].text = "Nombre"
t6.rows[0].cells[1].text = "Rol"
t6.rows[0].cells[2].text = "Fecha"
t6.rows[0].cells[3].text = "Conformidad"
t6.rows[1].cells[0].text = "[RELLENAR]"
t6.rows[1].cells[1].text = ""
t6.rows[1].cells[2].text = ""
t6.rows[1].cells[3].text = "OK / Pendiente"

add_p("Fin del Kickoff — siguiente: documentación AT01.", italic=True)

out = Path(r"C:\Users\HP\Documents\libreriaios\docs\00-KICKOFF.docx")
doc.save(out)
print(out)
print("bytes", out.stat().st_size)
